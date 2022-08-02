from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from functions.laplace_func import read_approx_file
from functions.equations_func import scien_format, fill_equation, latex_to_word
from functions.poles_table_func import get_poles_table



def sign_unicode(sign):
    if '+' in sign and '-' in sign:
        return u'\u00B1'
    elif '-' in sign:
        return u'\u002D'
    else:
        return ''
    
def head(d, device_type, freq_range):
    
    doc_style = d.styles['Normal']
    doc_style.font.name = 'Times New Roman'
    doc_style.font.size = Pt(14)


    header = d.sections[0].header.paragraphs[0]
    header.text = f'Model {device_type} {freq_range}'
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.style.font.bold = True
    header.style.font.size = Pt(16)    
    
def channel_info(d,channel_text, approx_file, math_equation, equation_legend, table_var):
    
    # Vars
    A0, zeros, poles = read_approx_file(approx_file)
    n_zeros, n_poles = len(zeros), len(poles)
    A0, pow_ten = scien_format(A0)
    poles_table = get_poles_table(poles)
    
    math_equation = fill_equation(math_equation, A0=A0, pow_ten=pow_ten,
                                  n_zeros=n_zeros, n_poles=n_poles)
    
    #Title
    
    text_channel = d.add_paragraph()
    text_channel = text_channel.add_run(channel_text)
    text_channel.bold = True
    text_channel.font.size = Pt(16)
    
    
    # Formula
    word_math = latex_to_word(math_equation)
    p = d.add_paragraph()
    p._element.append(word_math)
    
    d.add_paragraph(equation_legend)    
    
    # Table
    table1 = d.add_table(rows=len(poles_table), cols=3)
    
    for i, row in enumerate(table1.rows):
        poles_list = poles_table[i]['poles_list']
        poles_list = ', '.join(map(str, poles_list))
        sign = poles_table[i]['sign']
        sign = sign_unicode(sign)
        p1 = row.cells[0].paragraphs[0]
        p1_idx = p1.add_run(table_var)
        p1_idx.bold = True
        p1.add_run(poles_list).font.subscript = True
        row.cells[1].text = str(poles_table[i]['real'])
        row.cells[2].text = sign + str(poles_table[i]['imag'])
        
    d.add_paragraph()


    

approx_folder = '/home/lapa/Complex-curve-fit/RESP/examples/4211/!Device_No-387/'

device_name = '4211_50s-60Hz'
device_type = 'CME-' + device_name.split('_')[0]
freq_range = '(50s - 60Hz)'
sensitivity = 2000

channel_text = '{} channel ({} V/m/sec)'


d = Document()


head(d, device_type, freq_range)

for title in ['', 'Laplace']:
    
    if title == 'Laplace':
        math_equation = r'W(s) = \frac{A0 * 10^{pow_ten} * s^{n_zeros}}{\prod_{i=1}^{n_poles}{(s-s_{i})}}'
        equation_legend = 's - Laplace variable'
        table_var = 'S'
        prefix = 'Laplace_'
    else:
        math_equation = r'F(f) = \frac{A0 * 10^{pow_ten} * f^{n_zeros}}{\prod_{i=1}^{n_poles}{(f-p_{i})}}'
        equation_legend = 'f – frequency in Hz'
        table_var = 'P'
        prefix = ''
        
    tit = d.add_paragraph() 
    tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tit.add_run(title).bold = True

    
    
    for channel in ['Vert', 'Hor']:
        if channel == 'Vert':
            tmp_channel = channel_text.format('Vertical', sensitivity)
        else:
            tmp_channel = channel_text.format('Horizontal', sensitivity)
        
        approx_file = approx_folder + prefix + device_name + f'_{channel}'
        
        channel_info(d, tmp_channel, approx_file,
                     math_equation, equation_legend, table_var)
        

d.save(f'{approx_folder}{device_name}.docx')